#!/usr/bin/env python3
"""Builds the Project Meeting 3 report as a .docx.

Only the new and changed sections are produced. Unchanged material from the
Meeting 2 report (the full use case specifications, the appendix) is not
repeated here and stays in the existing document.
"""

import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = sys.argv[1] if len(sys.argv) > 1 else "ESC_Project_Meeting_3_Report.docx"

# Replace the three handles with the members' real names before submitting.
NAMES = {
    "systems": "Shadesofastar",
    "networking": "Wong Yee Jin",
    "application": "Jaelucard",
}

# Every figure here is taken from the repository, not estimated.
# git diff --shortstat $(git merge-base main Terminal-version) Terminal-version
STATS = {
    "commits": 34,
    "files_changed": 65,
    "insertions": "2,908",
    "deletions": "3,557",
    "unit_files": 12,
    "unit_cases": 79,        # make test: 79 Tests 0 Failures 0 Ignored
    "unit_assertions": 216,  # count of TEST_ASSERT* calls in tests/*.c
    "e2e_cases": 13,         # python3 -m unittest discover -s tests/e2e
}

BLUE = RGBColor(0x1F, 0x4E, 0x79)
LIGHT_BLUE = RGBColor(0x2E, 0x74, 0xB5)

doc = Document()

# Match the Meeting 2 document's look: blue headings, plain body text.
styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)

for name, size, colour in (
    ("Heading 1", 18, BLUE),
    ("Heading 2", 14, LIGHT_BLUE),
    ("Heading 3", 12, LIGHT_BLUE),
):
    st = styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = colour
    st.font.bold = False


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def p(text):
    return doc.add_paragraph(text)


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def code(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return para


def caption(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def figure_slot(label, filename):
    """Marks where a rendered PlantUML image is pasted in."""
    para = doc.add_paragraph()
    run = para.add_run(f"[ Paste rendered diagram here: {filename} ]")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(label)


# ---------------------------------------------------------------------------
# Title
# ---------------------------------------------------------------------------
title = doc.add_paragraph()
run = title.add_run("Mini GitHub Issue Tracker")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = BLUE
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
run = sub.add_run("Project Meeting 3 Report")
run.font.size = Pt(14)
run.font.color.rgb = LIGHT_BLUE
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub2 = doc.add_paragraph()
run = sub2.add_run("Team C3C8 Ranch Relishers")
run.font.size = Pt(11)
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER

p("This report covers the sections that changed since Project Meeting 2. The "
  "use case specifications and the build appendix from the previous report are "
  "unchanged except where stated in section 1 and are not reproduced here.")

# ---------------------------------------------------------------------------
# 1. Changes in Requirement
# ---------------------------------------------------------------------------
h1("1. Changes in Requirement")

p("There are substantial changes since Project Meeting 2. The largest is a "
  "change of delivery platform: the product is now a terminal application "
  "rather than a web application. This section separates the changes that came "
  "from tutor feedback from the ones the team decided on while carrying that "
  "feedback out.")

h2("1.1 Changes arising from Project Meeting 2 feedback")

p("Five points were raised at the previous meeting. Each is listed below with "
  "what was done about it.")

table(
    ["Feedback received", "Action taken", "Status"],
    [
        ["A separate web application with a UI is not required. Running the "
         "project on the terminal is sufficient.",
         "The web stack was removed and the product rebuilt as an interactive "
         "terminal application. The TLS server, HTTP parser, HTML templating "
         "and form handling were deleted, and the business logic was preserved "
         "behind a new service layer.",
         "Done"],
        ["The server must be running before the terminal application can "
         "start.",
         "Not yet implemented. The application currently opens its SQLite "
         "database directly in process. The client and server split is "
         "described in section 6 as remaining work.",
         "Outstanding"],
        ["Function names in the sequence diagrams must match the functions "
         "actually called in the code.",
         "All sequence diagrams were redrawn using the real C function names. "
         "Each diagram names the source file for every participant so a reader "
         "can check any message against the source.",
         "Done"],
        ["All arrows in the solution class diagram must be dotted.",
         "The solution class diagram was redrawn with every relationship as a "
         "dotted dependency arrow.",
         "Done"],
        ["The use cases are written properly and the diagrams are good.",
         "The use case specification format is unchanged. Two use cases were "
         "renamed to match the implemented behaviour, and eight were added. "
         "See section 1.2.",
         "Retained"],
    ],
    widths=[2.2, 3.0, 0.9],
)

h2("1.2 Changes the team made during the migration")

p("Moving off the web stack forced several decisions that change the "
  "requirements as written at Meeting 2.")

h3("Database changed from MongoDB to SQLite")

p("MongoDB requires a running server process, which conflicts with delivering "
  "a single terminal binary a marker can build and run. SQLite stores the "
  "whole database in one file and needs no server. The schema is created at "
  "startup with CREATE TABLE IF NOT EXISTS, so a first run on an empty machine "
  "works with no setup step.")

h3("Authentication changed from a web OAuth redirect to the GitHub device flow")

p("The previous design caught an OAuth redirect on a callback URL. A terminal "
  "application has no URL to redirect to. The device flow was designed for "
  "exactly this case: the application requests a code, the user enters that "
  "code in a browser, and the application polls GitHub until the "
  "authorisation completes. Credentials are read from the GH_CLIENT_ID "
  "environment variable, so no client identifier is compiled into the binary.")

h3("Server side sessions replaced by a local token file")

p("The web version stored session tokens in the database. The terminal version "
  "writes the access token to a file under the user configuration directory "
  "with permissions 0600. No token is stored in the database.")

h3("Time tracking removed")

p("Time estimation and time logging were dropped. The feature was already "
  "marked for removal in the Meeting 2 source, and the duration parsing code "
  "carried an integer overflow defect. Removing the feature removed the defect.")

h3("Two use cases renamed")

table(
    ["Meeting 2 name", "Meeting 3 name", "Reason"],
    [
        ["Sign Up",
         "Sign In with GitHub",
         "Registration and login are the same operation. The database records "
         "a user by GitHub identifier and either inserts a new row or updates "
         "the existing one, so there is no separate sign up path to describe."],
        ["Add Project Contributor",
         "Assign User to Issue",
         "Assignment is recorded per issue rather than per project, which "
         "matches how GitHub itself assigns work. There is no project "
         "membership concept in the system."],
    ],
    widths=[1.6, 1.8, 2.7],
)

h3("Search and filter narrowed to a single project")

p("At Meeting 2 search ran across all issues. It is now scoped to the project "
  "the user is viewing. This was done as a security fix. The previous "
  "implementation passed the user keyword into a MongoDB regular expression, "
  "which allowed a crafted keyword to change the query. The replacement binds "
  "the keyword as a SQL parameter, escapes the LIKE wildcards, restricts "
  "results to the current project and applies a row limit. Scoping is a real "
  "reduction in capability and is recorded here as a deliberate trade.")

h2("1.3 Use cases added since Meeting 2")

p("The ten use cases from Meeting 2 are retained, two under new names. Eight "
  "further use cases were implemented during this period.")

table(
    ["Use case", "Why it was added"],
    [
        ["Sign Out", "Clears the stored token and the signed in user, needed "
                     "once login was implemented."],
        ["Create Label", "Labels existed but could only be read. Creation "
                         "completes the feature."],
        ["Assign Label to Issue", "Categorisation was in the domain model at "
                                  "Meeting 2 but had no user path in the "
                                  "terminal version until now."],
        ["Add Comment to Issue", "Comments existed in the web version and were "
                                 "carried across."],
        ["Filter Issues", "Narrowing by status or by label, separate from "
                          "keyword search."],
        ["View My GitHub Repositories", "New capability. Once the device flow "
                                        "was working the same token could list "
                                        "the user's repositories."],
        ["Create Local User", "Without it the assignee list is empty unless the "
                              "user signs in to GitHub, which makes assignment "
                              "impossible offline."],
        ["View Assignees", "Lists the users known to the system."],
    ],
    widths=[2.0, 4.1],
)

# ---------------------------------------------------------------------------
# 2. Updated Design
# ---------------------------------------------------------------------------
doc.add_page_break()
h1("2. Updated Design")

p("All diagrams were redrawn for the terminal architecture. The source for "
  "each is committed under docs/diagrams in PlantUML format so the diagrams "
  "can be regenerated and reviewed alongside the code.")

h2("2.1 Use case diagram")

p("The primary actor is the User. GitHub Server remains a secondary actor but "
  "now participates in only two use cases, signing in and listing "
  "repositories. Every other use case runs with no network connection.")

figure_slot("Figure 1. Use case diagram for the terminal version.",
            "docs/diagrams/01-use-case.puml")

h2("2.2 Domain class diagram")

p("Three changes since Meeting 2. The Contributor entity was removed, because "
  "assignment is a relationship between Issue and User and does not need its "
  "own entity. Comment was promoted to a first class entity. The "
  "estimateMinutes and loggedMinutes attributes were removed from Issue along "
  "with the time tracking feature.")

figure_slot("Figure 2. Domain class diagram.",
            "docs/diagrams/02-domain-class.puml")

h2("2.3 Solution class diagram")

p("Every arrow is dotted, as requested at the previous meeting. Every "
  "operation carries the real C function name, and each box is annotated with "
  "the source file its functions live in.")

p("The structure is a four layer split. The boundary layer is the terminal "
  "interface. The control layer holds the business rules. The integration "
  "layer talks to GitHub. The data layer owns all SQL. C has no classes, so as "
  "in the Meeting 2 report each box is a translation unit together with the "
  "struct it operates on.")

p("One property of this design is worth stating because the tests depend on "
  "it. No function in the boundary layer calls the data layer directly. Every "
  "path from the user interface to the database passes through a service "
  "function, and every service function that writes calls "
  "auth_ctx_is_authed() before it reaches SQL. That is why the authorisation "
  "rule can be tested once per service rather than once per screen.")

figure_slot("Figure 3. Solution class diagram, all relationships dotted.",
            "docs/diagrams/03-solution-class.puml")

h2("2.4 Sequence diagrams")

p("Seven sequence diagrams cover the use cases that involve more than a single "
  "read. Message names are the C function names in the source.")

table(
    ["Figure", "Use case", "PlantUML source"],
    [
        ["Figure 4", "Sign In with GitHub", "04-seq-signin.puml"],
        ["Figure 5", "Create Project", "05-seq-create-project.puml"],
        ["Figure 6", "Create Issue", "06-seq-create-issue.puml"],
        ["Figure 7", "View Issue", "07-seq-view-issue.puml"],
        ["Figure 8", "Close Issue and Reopen Issue", "08-seq-close-reopen.puml"],
        ["Figure 9", "Search Issue", "09-seq-search-issue.puml"],
        ["Figure 10", "Assign User to Issue", "10-seq-assign-user.puml"],
    ],
    widths=[0.9, 2.6, 2.6],
)

p("One point needs stating because the diagram would otherwise look wrong "
  "against the use case list. Close Issue and Reopen Issue are two separate "
  "user goals but one code path. The web version had separate close and reopen "
  "routes. The terminal version has a single toggle whose target status is "
  "derived from the current status, so both use cases call "
  "issue_service_set_status with a different argument. Figure 8 shows both "
  "branches rather than pretending there are two functions.")

for label, fname in [
    ("Figure 4. Sign In with GitHub, device flow.", "04-seq-signin.puml"),
    ("Figure 5. Create Project.", "05-seq-create-project.puml"),
    ("Figure 6. Create Issue.", "06-seq-create-issue.puml"),
    ("Figure 7. View Issue.", "07-seq-view-issue.puml"),
    ("Figure 8. Close Issue and Reopen Issue.", "08-seq-close-reopen.puml"),
    ("Figure 9. Search Issue.", "09-seq-search-issue.puml"),
    ("Figure 10. Assign User to Issue.", "10-seq-assign-user.puml"),
]:
    figure_slot(label, "docs/diagrams/" + fname)

# ---------------------------------------------------------------------------
# 3. Implementation progress since Project Meeting 2
# ---------------------------------------------------------------------------
doc.add_page_break()
h1("3. Implementation Progress Since Project Meeting 2")

h2("3.1 Where the project stood at Meeting 2")

p("At the previous meeting the product was a web application written in C. It "
  "ran a TLS server on port 8443, forked a process per connection, parsed HTTP "
  "by hand, stored data in MongoDB, and generated HTML from C string builders "
  "in a 423 line template module. It exposed nineteen routes. It had no "
  "automated tests of any kind and no test target in the Makefile.")

h2("3.2 What was built since")

p("The work is recorded as " + str(STATS["commits"]) + " commits on the "
  "Terminal-version branch. Measured against the point where the branches "
  "diverged, " + str(STATS["files_changed"]) + " files changed, with "
  + str(STATS["insertions"]) + " insertions and " + str(STATS["deletions"])
  + " deletions.")

table(
    ["Area", "What was done", "Source"],
    [
        ["Persistence",
         "MongoDB replaced by SQLite. Schema created at startup. Every "
         "statement uses bound parameters. Per project issue numbering added.",
         "src/db.c"],
        ["Service layer",
         "Business rules extracted out of the HTTP handlers into a user "
         "interface agnostic layer. Authorisation enforced once per service "
         "rather than per route.",
         "src/core/"],
        ["Terminal interface",
         "New menu driven interface covering every screen the web version had.",
         "src/ui/menu.c"],
        ["Web stack removal",
         "TLS server, HTTP parser, HTML templating, form parsing and the web "
         "OAuth redirect deleted.",
         "commit f5c3e2e"],
        ["GitHub integration",
         "Device flow login, session resume, sign out, and repository listing "
         "over HTTPS with libcurl.",
         "src/github.c"],
        ["Token storage",
         "Access token written to a 0600 file under the user configuration "
         "directory instead of a database session table.",
         "src/token_store.c"],
        ["JSON parsing",
         "Hand written scanner replacing substring searching, so responses are "
         "parsed by structure rather than by text order.",
         "src/json.c"],
        ["Presentation",
         "Terminal size detection with three layout breakpoints, ANSI colour "
         "honouring NO_COLOR and non terminal output, and a skippable ASCII "
         "splash.",
         "src/render.c, src/assets.c"],
        ["Testing",
         "Test suite built from nothing. See section 4.",
         "tests/"],
    ],
    widths=[1.3, 3.5, 1.3],
)

h2("3.3 Use case implementation status")

p("Every functional use case is implemented. Two are marked as verified "
  "manually rather than automatically, because they require a live GitHub "
  "account and a browser step that cannot run unattended.")

table(
    ["#", "Use case", "Status", "Implementing function"],
    [
        ["1", "Sign In with GitHub", "Implemented, manually verified",
         "github_login, github_device_start, github_device_poll"],
        ["2", "View Projects", "Fully implemented",
         "screen_projects, project_service_list"],
        ["3", "Create Project", "Fully implemented",
         "project_service_create"],
        ["4", "View Issue", "Fully implemented",
         "screen_issue_detail, issue_service_get"],
        ["5", "Create Issue", "Fully implemented", "issue_service_create"],
        ["6", "Close Issue", "Fully implemented",
         "issue_service_set_status(STATUS_CLOSED)"],
        ["7", "Reopen Issue", "Fully implemented",
         "issue_service_set_status(STATUS_OPEN)"],
        ["8", "View Labels", "Fully implemented", "label_service_list"],
        ["9", "Search Issue", "Fully implemented",
         "issue_service_search, db_issue_search"],
        ["10", "Assign User to Issue", "Fully implemented",
         "issue_service_add_assignee"],
        ["11", "Sign Out", "Fully implemented", "github_logout, token_clear"],
        ["12", "Create Label", "Fully implemented", "label_service_create"],
        ["13", "Assign Label to Issue", "Fully implemented",
         "issue_service_add_label"],
        ["14", "Add Comment to Issue", "Fully implemented",
         "comment_service_add"],
        ["15", "Filter Issues", "Implemented with a usability defect",
         "issue_service_filter"],
        ["16", "View My GitHub Repositories", "Implemented, manually verified",
         "github_repos, github_list_repos"],
        ["17", "Create Local User", "Fully implemented", "user_service_create"],
        ["18", "View Assignees", "Fully implemented",
         "screen_assignees, user_service_list"],
    ],
    widths=[0.3, 1.6, 1.6, 2.6],
)

p("No use case is unimplemented. The defect against use case 15 is described "
  "in section 6.")

# ---------------------------------------------------------------------------
# 4. Testing
# ---------------------------------------------------------------------------
doc.add_page_break()
h1("4. Testing")

p("At Meeting 2 the project had no tests. It now has two suites, both run from "
  "the Makefile.")

h2("4.1 Unit tests")

p("Unit tests are written with Unity 2.7.1, an established C unit testing "
  "framework. "
  "C has no equivalent of Jest or RSpec in the sense of a single dominant "
  "tool, and Unity is the closest counterpart: it provides named test cases, "
  "per test setUp and tearDown, typed assertion macros that report expected "
  "against actual values, and a runner that reports pass and fail counts. It "
  "is vendored into the repository under tests/vendor/unity, so no package "
  "needs to be installed to build or run the suite.")

p("There are " + str(STATS["unit_cases"]) + " test cases across "
  + str(STATS["unit_files"]) + " modules, containing "
  + str(STATS["unit_assertions"]) + " assertions, up from 150 assertions "
  "before the framework was adopted. All pass. No existing assertion was "
  "removed or weakened during the move: every file has at least as many "
  "assertions as it started with.")

p("Tests that touch the database run against an in memory SQLite instance "
  "created in setUp and destroyed in tearDown. This was a real improvement "
  "rather than a tidy up. The previous versions opened one database at the top "
  "of the file and shared it across every assertion in that file, so a test "
  "could pass because of a record an earlier test had left behind, and the "
  "order the assertions were written in silently mattered. The signed in user "
  "is reset the same way, so an authorisation check cannot pass on state "
  "leaked from the test above it.")

p("Test case names follow the use cases they exercise, so the suite can be "
  "read against the requirements. Coverage by module:")

table(
    ["Module", "What the tests verify"],
    [
        ["db.c",
         "Create, find and list for every entity. Per project issue number "
         "allocation. Label and assignee join tables including repeat "
         "assignment. Project scoped search proven with a decoy issue in a "
         "second project. Wildcard characters in a keyword treated literally "
         "rather than as patterns. Null keyword and title handled without "
         "crashing."],
        ["core/issue_service.c",
         "Authorisation refused when signed out and granted when signed in, "
         "for create, status change, labelling and assignment. Title "
         "validation. Search and filter scoping."],
        ["core/project_service.c",
         "Name uniqueness, blank name rejection, authorisation gate."],
        ["core/label_service.c", "Creation, uniqueness, assignment to an issue."],
        ["core/user_service.c",
         "GitHub upsert taking the insert path on first sign in and the update "
         "path on return, confirming the second sign in reuses the same "
         "record. Local user creation refused when signed out, rejected for a "
         "blank name, rejected for a duplicate, and stored and readable back "
         "when valid."],
        ["core/comment_service.c", "Adding a comment, parent issue existence "
                                   "check, listing in insertion order."],
        ["github.c",
         "Device code, access token and repository list responses parsed from "
         "captured sample JSON. Pending, slow down, expired and denied error "
         "responses mapped to the right status."],
        ["json.c",
         "Field extraction, integer fields, arrays of objects, and a nested "
         "key miss that proves the parser reads structure rather than "
         "searching for text."],
        ["render.c",
         "Breakpoint selection for given terminal dimensions. Colour enabled "
         "and disabled correctly for NO_COLOR and for non terminal output."],
        ["token_store.c",
         "Save and load round trip. File permissions asserted to be exactly "
         "0600. Behaviour when no token is stored."],
        ["util.c", "Identifier length and uniqueness, and handling of a "
                   "failure from the random number source."],
        ["ui/menu.c", "Input trimming and menu choice parsing, including "
                      "non numeric input and an empty line."],
    ],
    widths=[1.4, 4.7],
)

h2("4.2 End to end tests")

p("Unit tests call the service functions directly, which means they cannot "
  "catch a fault in the interface itself. A screen that never calls its "
  "service, a menu option wired to the wrong action, or a result that is "
  "computed correctly and then printed wrongly would all pass every unit test. "
  "The whole of the terminal interface sits below that line.")

p("The second suite closes that gap by driving the compiled binary from "
  "outside. It is written with Python's unittest framework and uses no third "
  "party packages. Each test starts ./mini-gh-tracker as a subprocess with "
  "DB_PATH pointed at a fresh temporary database, writes keystrokes to its "
  "standard input, and asserts on the text the application renders. This is "
  "the same approach Selenium or Cypress take with a browser, applied to a "
  "terminal, which is this product's user interface.")

p("There are " + str(STATS["e2e_cases"]) + " scenarios. They cover creating "
  "and listing projects, creating issues and checking per project numbering, "
  "opening an issue and reading its rendered detail, closing and reopening "
  "through the status toggle, listing and creating labels, assigning a user, "
  "keyword search returning matches and excluding non matches, empty state "
  "messages, and recovery from invalid menu input.")

p("Twelve of the thirteen pass. The thirteenth, "
  "test_filter_by_label_name_is_broken, is marked as an expected failure. It "
  "is written the way the feature ought to behave, filtering by a label name "
  "the user can see, and it fails because the current screen asks for an "
  "internal identifier instead. Recording the defect as a failing test rather "
  "than a comment means the suite will tell us when it is fixed: once the "
  "behaviour is corrected the test starts passing, and an unexpected pass is "
  "reported rather than silently ignored.")

h2("4.3 What is not covered")

table(
    ["Area", "Reason", "How it is verified instead"],
    [
        ["Live GitHub device flow",
         "Requires a real account and a manual browser authorisation step, so "
         "it cannot run unattended.",
         "Response parsing is unit tested against captured JSON. The flow "
         "itself is checked against a written manual checklist."],
        ["Terminal size detection and the splash animation",
         "Require a real terminal device. Piped output is not a terminal.",
         "The decision functions that choose a layout mode and enable colour "
         "are pure and are unit tested directly."],
    ],
    widths=[1.6, 2.4, 2.5],
)

h2("4.4 Running the tests")

code("make check     # both suites\n"
     "make test      # unit tests only\n"
     "make e2e       # end to end tests only")

p("Neither suite needs anything installed. Unity is vendored in the "
  "repository and the end to end suite uses only the Python standard library.")

h2("4.5 Test isolation")

p("Both suites are hermetic, which is worth stating because the first version "
  "of the end to end suite was not. It set the database path to a temporary "
  "file but inherited the rest of the environment, so the application found "
  "the developer's own cached GitHub token at startup and tried to revalidate "
  "it over the network. On a machine where nobody had signed in the tests "
  "passed and the fault was invisible. On a machine where somebody had, every "
  "test would have made a live API call and timed out.")

p("Both suites now redirect the configuration directory at a scratch path, so "
  "no test reads or overwrites a real token and no test reaches the network. "
  "This was confirmed by planting a token in the environment and rerunning the "
  "suite, which produced identical results in under a tenth of a second.")

# ---------------------------------------------------------------------------
# 5. Feature progress and workload distribution
# ---------------------------------------------------------------------------
doc.add_page_break()
h1("5. Feature Progress and Workload Distribution")

h2("5.1 Roles")

table(
    ["Member", "Role", "Area of responsibility"],
    [
        [NAMES["systems"], "Systems",
         "Build, data layer, application core, platform migration."],
        [NAMES["networking"], "Networking and Security",
         "Transport, authentication, GitHub integration, credential handling."],
        [NAMES["application"], "Application and Integration",
         "Interface behaviour, feature integration, test suites."],
    ],
    widths=[1.5, 1.7, 2.9],
)

h2("5.2 Contribution up to Project Meeting 2")

p("The web application was built with work split across the team. Figures are "
  "taken from the commit history up to the point where the terminal branch "
  "diverged.")

table(
    ["Member", "Modules owned", "Lines added"],
    [
        [NAMES["application"],
         "db.c, template.c, user_handlers.c, router.c and their headers",
         "1,538"],
        [NAMES["networking"],
         "oauth_github.c, htttp.c, auth_handlers.c, main.c, secure_session.c, "
         "auth.c and their headers",
         "1,313"],
        [NAMES["systems"],
         "issue_handlers.c, time_handlers.c, label_handlers.c, "
         "project_handlers.c, comment_handlers.c, models.h, form_util.c",
         "867"],
        [NAMES["networking"] + " (repository setup)",
         "Certificate generation, README, licence",
         "71"],
    ],
    widths=[1.5, 3.4, 1.2],
)

h2("5.3 Contribution between Meeting 2 and Meeting 3")

p("This period is not evenly distributed and the report states that plainly. "
  "The platform migration was carried out by the Systems member in a "
  "concentrated period of work on 25 July 2026. All "
  + str(STATS["commits"]) + " commits on the Terminal-version branch since the "
  "branches diverged are under that member's authorship.")

table(
    ["Member", "Work in this period", "Commits"],
    [
        [NAMES["systems"],
         "Complete platform migration: SQLite data layer, core service layer, "
         "terminal interface, web stack removal, GitHub device flow, token "
         "storage, JSON parser, presentation layer, and the initial test "
         "suite.",
         str(STATS["commits"])],
        [NAMES["networking"],
         "Continued work on the web branch during the same period, including "
         "automatic project population on sign in (commit 7ccac7c on main). "
         "That work has not been carried across to the terminal branch.",
         "1 (on main)"],
        [NAMES["application"],
         "Test ownership is being handed over during this period. The suites "
         "described in section 4 are prepared for this member to take on and "
         "extend.",
         "Handover in progress"],
    ],
    widths=[1.5, 3.4, 1.2],
)

p("Two consequences follow and both are being addressed. The first is that the "
  "team worked in two directions at once, with the migration on one branch and "
  "continued web development on another. A single branch has to be chosen and "
  "the outstanding web commit either ported or dropped. The second is that "
  "test ownership sits with the Application and Integration member by "
  "agreement, and the handover of the two suites is the immediate next step "
  "for that member.")

h2("5.4 Deliverables by module")

p("Each claim above can be checked against a specific file.")

table(
    ["Deliverable", "Location"],
    [
        ["SQLite data layer with parameterised queries", "src/db.c, include/db.h"],
        ["Core service layer and authorisation gate",
         "src/core/, include/core/services.h, include/core/auth_ctx.h"],
        ["Terminal interface", "src/ui/menu.c, include/ui/menu.h"],
        ["GitHub device flow and repository listing", "src/github.c, include/github.h"],
        ["Token storage", "src/token_store.c"],
        ["JSON parser", "src/json.c"],
        ["Presentation and layout", "src/render.c, src/assets.c"],
        ["Unit test suite", "tests/test_*.c, tests/vendor/unity/"],
        ["End to end test suite", "tests/e2e/"],
        ["Design specification for the migration",
         "docs/superpowers/specs/2026-07-25-terminal-migration-design.md"],
        ["Diagram sources", "docs/diagrams/*.puml"],
    ],
    widths=[3.0, 3.1],
)

# ---------------------------------------------------------------------------
# 6. Known defects and remaining work
# ---------------------------------------------------------------------------
doc.add_page_break()
h1("6. Known Defects and Remaining Work")

h2("6.1 Known defects")

p("These were found by reviewing the code against the design specification "
  "rather than by a user reporting them. None of them stops the product "
  "working, and each is listed with where it lives.")

table(
    ["Defect", "Effect", "Location"],
    [
        ["Filtering by label prompts for a raw 24 character identifier",
         "The user has to type an internal identifier by hand, while the "
         "neighbouring add label screen offers a numbered list. The feature "
         "works but is not usable in practice.",
         "src/ui/menu.c, filter branch"],
        ["Comments record no author",
         "A comment is stored with its text and time but no user. The "
         "authorisation gate confirms someone is signed in and then discards "
         "which user it was. The comments table has no column for it.",
         "src/db.c schema, src/core/comment_service.c"],
        ["Stale token handling is inconsistent",
         "If the network fails while a stored token is being revalidated at "
         "startup, the menu offers a full sign in again rather than retrying "
         "the cached token, even though the repository listing would still "
         "work with it.",
         "src/ui/menu.c, session resume"],
        ["Terminal resize is not handled live",
         "The design specified a signal handler so the display re-lays out "
         "immediately. It was not implemented. The effect is limited because "
         "every screen re-queries the terminal size when it draws, so a "
         "resize takes effect on the next screen rather than at once.",
         "src/render.c"],
        ["Two service functions are unreachable from the interface",
         "project_service_get and user_service_get are implemented and tested "
         "but nothing in the interface calls them.",
         "src/core/"],
    ],
    widths=[1.7, 3.0, 1.4],
)

h2("6.2 Remaining work")

h3("Separating the application from the server")

p("This is the outstanding item from Meeting 2 feedback. The application "
  "currently opens the SQLite database directly in its own process, at "
  "src/main.c, before the menu loop starts. There is no server process and no "
  "socket. The feedback asked for a model where the server must already be "
  "running before the terminal application starts.")

p("The design is already shaped for this change. No interface code touches the "
  "database. Every call goes through the service layer, so the split can be "
  "made at that boundary: the service functions become client stubs that send "
  "a request to the server process, and the existing implementations move "
  "behind the server. The planned steps are to define the request and response "
  "format, host the existing service and data layers in a server process that "
  "owns the database, replace the service layer in the client with calls over "
  "a local socket, and have the client fail with a clear message when the "
  "server is not running.")

h3("Other planned work")

table(
    ["Item", "Description", "Owner"],
    [
        ["Fix the label filter",
         "Present a numbered list of labels, matching the add label screen, "
         "instead of asking for an identifier.",
         "Application and Integration"],
        ["Record comment authors",
         "Add a user column to the comments table and populate it from the "
         "signed in user, which is already tracked but not used.",
         "Systems"],
        ["Extend end to end coverage",
         "Add scenarios for the GitHub screens using a stub, and for comment "
         "and label assignment paths.",
         "Application and Integration"],
        ["Resolve the branch split",
         "Decide whether the outstanding web branch commit is ported to the "
         "terminal version or dropped, and settle on a single branch.",
         "Whole team"],
        ["Handle terminal resize immediately",
         "Add the signal handler described in the design specification.",
         "Systems"],
    ],
    widths=[1.4, 3.2, 1.5],
)

doc.save(OUT)
print(f"wrote {OUT}")
